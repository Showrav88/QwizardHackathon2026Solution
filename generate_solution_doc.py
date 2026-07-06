#!/usr/bin/env python3
"""Generate the full non-AI contact center solution Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    return doc.add_paragraph(text, style="List Bullet")


def add_table(doc, headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_color)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # Title page
    title = doc.add_heading("Contact Center Optimization Solution", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Non-AI Implementation Plan\n"
        "Addressing the Qwizard Hackathon 2026 Challenge Without Machine Learning"
    )
    run.font.size = Pt(14)
    run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\n\nPrepared for: Your Company\n")
    meta.add_run("Document Version: 1.0\n")
    meta.add_run("Date: July 2026\n")
    meta.add_run("Classification: Internal — Operational Improvement\n")

    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Problem Statement & Business Impact",
        "3. Hackathon Challenge Reference (AI Solution)",
        "4. Non-AI Solution Overview",
        "5. Component A: Agent Operations Toolkit",
        "6. Component B: Supervisor Command Center",
        "7. Component C: Quality Assurance Program",
        "8. Component D: Customer Experience & Call Deflection",
        "9. Technology Architecture (Non-AI Stack)",
        "10. 90-Day Implementation Roadmap",
        "11. KPIs, Targets & Measurement Framework",
        "12. Governance & Human-in-the-Loop Controls",
        "13. Risk Assessment & Mitigation",
        "14. Budget & Resource Estimates",
        "15. Appendices",
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This document presents a complete, production-ready plan to solve the same contact center "
        "problems targeted by the Qwizard Hackathon 2026 \"Real-Time AI Co-Pilot Challenge\" — "
        "without deploying artificial intelligence, machine learning, or real-time speech analytics."
    )
    add_para(
        doc,
        "The underlying business pains are process gaps, knowledge findability failures, and "
        "supervisory blind spots during live calls. These are well-understood operational problems "
        "with proven non-AI remedies: structured playbooks, CTI screen-pop, rule-based alerting, "
        "native PBX coaching features, and systematic QA sampling."
    )
    add_para(doc, "Expected outcomes within 90 days:", bold=True)
    outcomes = [
        "15–25% reduction in Average Handle Time (AHT)",
        "10–20% improvement in First Call Resolution (FCR)",
        "100% of agents reviewed weekly via structured QA sampling",
        "30–40% faster new-agent ramp-up through standardized playbooks",
        "Real-time supervisor visibility and intervention on every escalated call",
    ]
    for o in outcomes:
        add_bullet(doc, o)

    add_para(
        doc,
        "Total estimated investment: Low to moderate — primarily internal labor, existing telephony "
        "features, and lightweight tooling (wiki, wallboard, scorecard spreadsheet or simple web app). "
        "No AI licensing, GPU infrastructure, or speech-to-text pipelines required."
    )

    # 2. Problem Statement
    add_heading(doc, "2. Problem Statement & Business Impact", 1)
    add_heading(doc, "2.1 Core Problem", 2)
    add_para(
        doc,
        "The contact center operates without real-time operational support. Agents handle live calls "
        "in isolation, supervisors cannot see which interactions are deteriorating until it is too late, "
        "and quality assurance covers only a tiny fraction of total call volume. Customers experience "
        "long waits, repeated information requests, inconsistent answers, and unresolved issues."
    )

    add_heading(doc, "2.2 Affected Personas", 2)
    add_table(
        doc,
        ["Persona", "Current Experience", "Primary Pain", "Business Impact"],
        [
            [
                "Customer",
                "Long queue waits; must repeat account details; receives inconsistent answers",
                "Poor service; unresolved issues; repeat callbacks",
                "Churn, negative CSAT, increased call volume",
            ],
            [
                "Call Center Agent",
                "Handles live calls alone; searches knowledge bases mid-call; follows scripts from memory",
                "Slow resolution; high stress; compliance errors",
                "High AHT, low FCR, agent turnover",
            ],
            [
                "Supervisor / Team Lead",
                "No live visibility into call quality or escalation risk",
                "Intervention happens too late or not at all",
                "Escalations reach management; customer complaints",
            ],
            [
                "QA Reviewer",
                "Manually reviews 2–5% of recorded calls after the fact",
                "Most quality issues never detected",
                "Compliance risk; coaching gaps; no trend visibility",
            ],
        ],
    )

    add_heading(doc, "2.3 Target Business Metrics", 2)
    add_table(
        doc,
        ["Metric", "Definition", "Current State (Typical)", "Target Improvement"],
        [
            ["Average Handle Time (AHT)", "Total talk + hold + wrap-up time per call", "8–12 minutes", "15–25% reduction"],
            ["First Call Resolution (FCR)", "% of issues resolved without callback", "65–75%", "10–20% improvement"],
            ["QA Coverage", "% of agents reviewed per period", "2–5% of calls", "100% of agents weekly"],
            ["Agent Ramp-Up Time", "Days until new hire meets productivity benchmark", "60–90 days", "30–40% faster"],
            ["Queue Abandonment Rate", "% of callers who hang up before answer", "5–12%", "30–50% reduction"],
            ["Customer Satisfaction (CSAT)", "Post-call survey score", "Baseline TBD", "10+ point improvement"],
        ],
    )

    # 3. Hackathon Reference
    add_heading(doc, "3. Hackathon Challenge Reference (AI Solution)", 1)
    add_para(
        doc,
        "The Qwizard Hackathon 2026 July episode — \"Real-Time AI Co-Pilot Challenge\" — defines the "
        "problem framing and the shape of an AI-powered solution. The hackathon template is a blank "
        "submission form; it does not contain a completed team solution. It mandates two components:"
    )

    add_heading(doc, "3.1 Mandatory Component 1: Agent Co-Pilot", 2)
    features = [
        "Real-time audio ingestion (SIP/RTP or WebRTC)",
        "Live speech transcription with PII masking",
        "Customer sentiment and intent detection",
        "Knowledge article and next-best-action suggestions pushed to agent screen mid-call",
        "Agentic cognitive loop: Observe → Decide → Act → Adapt",
        "Human-in-the-loop governance (AI assists; agent retains authority)",
    ]
    for f in features:
        add_bullet(doc, f)

    add_heading(doc, "3.2 Mandatory Component 2: Supervisor Dashboard", 2)
    features = [
        "Real-time view of all active calls with sentiment/escalation alerts",
        "Identification of worst-performing calls on the floor",
        "Intervention capabilities: whisper, listen, barge, or take over",
        "Live metrics: queue depth, wait times, agent availability",
    ]
    for f in features:
        add_bullet(doc, f)

    add_heading(doc, "3.3 AI-to-Non-AI Mapping", 2)
    add_table(
        doc,
        ["Hackathon AI Feature", "Non-AI Equivalent", "Implementation Complexity"],
        [
            ["Real-time transcription", "Mandatory wrap-up disposition codes + post-call survey", "Low"],
            ["Sentiment detection", "Rule-based alerts (duration, transfers, hold time)", "Low"],
            ["Intent classification", "IVR routing + skills-based assignment + disposition tagging", "Medium"],
            ["Knowledge suggestions", "Searchable playbooks + CRM decision trees", "Low–Medium"],
            ["Next-best-action prompts", "Call-type playbooks with step-by-step workflows", "Low"],
            ["Supervisor escalation alerts", "Agent help button + wallboard rule alerts", "Low"],
            ["Live coaching", "Native PBX whisper/barge features", "Low (already included)"],
            ["100% call analysis", "Structured random sampling + post-call IVR survey", "Medium"],
        ],
    )

    # 4. Non-AI Solution Overview
    add_heading(doc, "4. Non-AI Solution Overview", 1)
    add_para(
        doc,
        "The non-AI solution replaces AI inference with deterministic process, visibility, and "
        "knowledge management. It delivers the same operational outcomes through four integrated pillars:"
    )
    add_table(
        doc,
        ["Pillar", "Replaces", "Key Deliverables"],
        [
            ["A. Agent Operations Toolkit", "AI Co-Pilot", "Playbooks, CTI screen-pop, canned responses, escalate button"],
            ["B. Supervisor Command Center", "AI Dashboard", "Real-time wallboard, rule alerts, whisper/barge coaching"],
            ["C. Quality Assurance Program", "AI transcription & scoring", "Scorecard, weekly sampling, disposition reporting"],
            ["D. Customer Experience Layer", "Proactive AI outreach", "IVR optimization, self-service, callback option"],
        ],
    )

    add_para(
        doc,
        "Design principle: Every recommendation is actionable by a human agent or supervisor. "
        "No black-box predictions. All alerts are explainable rule triggers. All knowledge is "
        "curated and approved by subject-matter experts."
    )

    # 5. Component A
    add_heading(doc, "5. Component A: Agent Operations Toolkit", 1)
    add_heading(doc, "5.1 Internal Knowledge Base & Call Playbooks", 2)
    add_para(
        doc,
        "Build a fast, searchable internal knowledge base. The primary agent delay during calls "
        "is answer findability — not a lack of intelligence. Fixing search and structure resolves "
        "most mid-call hesitation."
    )
    add_para(doc, "Playbook structure (one page per call type):", bold=True)
    playbook_sections = [
        "Call type name and trigger phrases (what the customer says)",
        "Required identity verification steps",
        "Step-by-step resolution workflow (numbered, no ambiguity)",
        "Approved wording for sensitive statements (refunds, cancellations, compliance)",
        "Escalation criteria (when to transfer or request supervisor)",
        "Required wrap-up disposition code",
        "Average target handle time for this call type",
    ]
    for s in playbook_sections:
        add_bullet(doc, s)

    add_para(doc, "Priority: Create playbooks for the top 20 call types within Week 1–2.", bold=True)

    add_heading(doc, "5.2 Screen-Pop CTI Integration", 2)
    add_para(
        doc,
        "When a call arrives, the phone system automatically opens the customer's CRM record, "
        "recent ticket history, and last disposition. The caller should never need to repeat "
        "account information."
    )
    add_para(doc, "Implementation requirements:", bold=True)
    add_bullet(doc, "PBX webhook or CTI connector on inbound call event")
    add_bullet(doc, "CRM lookup by caller ID (ANI) or IVR-entered account number")
    add_bullet(doc, "Auto-populate agent desktop with: customer name, account status, open tickets, last 3 interactions")
    add_bullet(doc, "Supported platforms: Asterisk/FreePBX (AMI), 3CX (API), Genesys (CTI), Twilio (webhook), any system with REST/webhook capability")

    add_heading(doc, "5.3 Canned Responses & Disposition Macros", 2)
    add_para(
        doc,
        "Pre-approved text snippets and one-click disposition macros for the top 20 call reasons. "
        "Agents select the macro; the system logs the disposition and can trigger follow-up workflows."
    )
    add_table(
        doc,
        ["Call Reason", "Macro Name", "Disposition Code", "Follow-Up Action"],
        [
            ["Billing error inquiry", "BILL-ERROR-VERIFY", "BILL_ERR", "Create billing adjustment ticket"],
            ["Refund request", "REFUND-ELIGIBILITY", "REFUND", "Check refund policy; process or escalate"],
            ["Account cancellation", "CANCEL-RETENTION", "CANCEL", "Offer retention; process if confirmed"],
            ["Password reset", "PWD-RESET-VERIFY", "PWD_RST", "Send reset link after verification"],
            ["Service outage report", "OUTAGE-CHECK", "OUTAGE", "Check status page; create incident ticket"],
            ["Shipping delay", "SHIP-TRACK", "SHIP_DLY", "Pull tracking; offer compensation if applicable"],
            ["Product return", "RETURN-RMA", "RETURN", "Issue RMA number; explain return policy"],
            ["Upgrade inquiry", "UPGRADE-OPTIONS", "UPGRADE", "Present tier options; transfer to sales if needed"],
            ["Complaint — agent behavior", "COMP-AGENT", "COMP_AGT", "Log complaint; supervisor callback within 24h"],
            ["Technical support — Tier 1", "TECH-T1-TRIAGE", "TECH_T1", "Run diagnostic checklist; escalate to T2 if needed"],
        ],
    )

    add_heading(doc, "5.4 One-Click Supervisor Escalation", 2)
    add_para(
        doc,
        "A dedicated \"Request Help\" button on the agent desktop. When clicked:"
    )
    steps = [
        "Supervisor receives instant notification (desktop alert, SMS, or team chat)",
        "Notification includes: agent name, customer ID, call duration, disposition so far, free-text reason",
        "Supervisor can whisper-coach (customer cannot hear) or barge (join the call)",
        "Event is logged for QA and coaching follow-up",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"Step {i}: {s}")

    # 6. Component B
    add_heading(doc, "6. Component B: Supervisor Command Center", 1)
    add_heading(doc, "6.1 Real-Time Wallboard", 2)
    add_para(
        doc,
        "Deploy a real-time wallboard using native telephony platform metrics. No custom AI required."
    )
    add_table(
        doc,
        ["Metric", "Source", "Display", "Refresh Rate"],
        [
            ["Calls in queue", "PBX ACD stats", "Large counter with color threshold", "5 seconds"],
            ["Longest wait time", "PBX queue monitor", "Timer with red alert > 3 min", "5 seconds"],
            ["Active call count", "PBX channel stats", "Per-team breakdown", "5 seconds"],
            ["Agents available / on call / wrap-up", "PBX agent state", "Status grid", "5 seconds"],
            ["Abandon rate (today)", "PBX CDR aggregation", "Percentage with trend arrow", "60 seconds"],
            ["Average handle time (today)", "PBX CDR aggregation", "Minutes:seconds vs. target", "60 seconds"],
        ],
    )

    add_heading(doc, "6.2 Rule-Based Alert Engine", 2)
    add_para(
        doc,
        "Deterministic alerts replace AI sentiment detection. Each rule is transparent and auditable."
    )
    add_table(
        doc,
        ["Alert Rule", "Trigger Condition", "Severity", "Supervisor Action"],
        [
            ["Long call", "Call duration > 12 minutes", "Warning", "Monitor; prepare to whisper-coach"],
            ["Extended hold", "Customer on hold > 5 minutes", "Warning", "Check agent status"],
            ["Repeat transfer", "2nd transfer on same call", "High", "Supervisor joins or takes over"],
            ["Queue backup", "Queue wait > 3 minutes", "High", "Reassign agents; trigger callback offer"],
            ["Agent help request", "Agent clicks \"Request Help\"", "Immediate", "Supervisor responds within 60 seconds"],
            ["Wrap-up exceeded", "Agent in wrap-up > 3 minutes", "Info", "Check if agent needs assistance"],
            ["Abandon spike", "> 3 abandons in 15 minutes", "High", "Investigate queue/routing issue"],
        ],
    )

    add_heading(doc, "6.3 Native PBX Coaching Features", 2)
    add_para(
        doc,
        "Nearly every modern PBX ships with supervisor intervention tools. Enable and train on:"
    )
    add_bullet(doc, "Listen — supervisor hears both parties; neither party knows")
    add_bullet(doc, "Whisper — supervisor speaks to agent only; customer cannot hear")
    add_bullet(doc, "Barge — supervisor joins the call; all parties can hear")
    add_bullet(doc, "Takeover — supervisor assumes the call from the agent")

  # 7. Component C
    add_heading(doc, "7. Component C: Quality Assurance Program", 1)
    add_heading(doc, "7.1 Standardized QA Scorecard", 2)
    add_para(
        doc,
        "A consistent scorecard applied to every reviewed call. Score each section 1–5; "
        "total score out of 100."
    )
    add_table(
        doc,
        ["Section", "Criteria", "Max Points", "Fail Condition"],
        [
            ["Greeting & Identification", "Agent states name, company; verifies caller identity", "15", "No verification on account change"],
            ["Active Listening", "Agent acknowledges concern; no unnecessary interruptions", "15", "Interrupts customer 3+ times"],
            ["Problem Diagnosis", "Agent asks clarifying questions; confirms understanding", "15", "Proceeds without understanding issue"],
            ["Resolution Steps", "Follows playbook; completes all required steps", "25", "Skips mandatory compliance step"],
            ["Compliance Statements", "Required disclosures delivered (varies by call type)", "15", "Missing legally required disclosure"],
            ["Closing & Confirmation", "Summarizes resolution; confirms satisfaction; offers additional help", "15", "Ends call without resolution summary"],
        ],
    )

    add_heading(doc, "7.2 Structured Random Sampling", 2)
    add_para(
        doc,
        "Review 5 calls per agent per week, selected randomly from the week's completed calls. "
        "This ensures 100% agent coverage weekly while keeping review workload manageable."
    )
    add_para(doc, "Sampling formula:", bold=True)
    add_para(doc, "Weekly QA hours needed = (Number of agents × 5 calls × 8 min review) / 60")
    add_para(doc, "Example: 50 agents = 250 calls/week = ~33 QA hours/week (less than 1 FTE).")

    add_heading(doc, "7.3 Post-Call Customer Survey", 2)
    add_para(
        doc,
        "Deploy a 1-question post-call survey on every completed call via IVR or SMS:"
    )
    add_bullet(doc, "\"Was your issue resolved today?\" — Yes / No / Partially")
    add_bullet(doc, "Optional second question: \"Rate your experience 1–5\"")
    add_para(
        doc,
        "This provides FCR and CSAT data on 100% of surveyed calls — broader coverage than "
        "any manual QA program and a direct customer-voice metric."
    )

    add_heading(doc, "7.4 Disposition Reporting & Trend Analysis", 2)
    add_para(
        doc,
        "Mandatory wrap-up codes on every call enable weekly reporting:"
    )
    add_bullet(doc, "Top 10 call drivers by volume (feeds playbook prioritization)")
    add_bullet(doc, "FCR rate by disposition code")
    add_bullet(doc, "AHT by disposition code (identifies inefficient call types)")
    add_bullet(doc, "Escalation rate by agent and team")
    add_bullet(doc, "QA score trend by agent (coaching trigger when score drops below 70)")

    # 8. Component D
    add_heading(doc, "8. Component D: Customer Experience & Call Deflection", 1)
    add_heading(doc, "8.1 IVR & Skills-Based Routing Optimization", 2)
    add_para(
        doc,
        "The single highest-ROI FCR improvement: route calls to the right agent the first time."
    )
    add_bullet(doc, "Audit current IVR menu against top 20 disposition codes")
    add_bullet(doc, "Add skills-based routing: billing agents get billing calls, tech agents get tech calls")
    add_bullet(doc, "Implement priority routing for VIP/enterprise accounts")
    add_bullet(doc, "Offer callback option when queue wait exceeds 3 minutes")

    add_heading(doc, "8.2 Self-Service Deflection", 2)
    add_para(
        doc,
        "Publish self-service pages and IVR options for the top 5 call drivers:"
    )
    add_table(
        doc,
        ["Call Driver", "Self-Service Option", "Expected Deflection Rate"],
        [
            ["Password reset", "Automated IVR reset after verification", "60–80%"],
            ["Order status / tracking", "Web portal + IVR tracking lookup", "40–60%"],
            ["Billing statement copy", "Email/PDF download from account portal", "50–70%"],
            ["FAQ — product how-to", "Searchable help center with top 20 articles", "20–30%"],
            ["Account balance inquiry", "IVR balance readout after verification", "40–50%"],
        ],
    )

    # 9. Technology Architecture
    add_heading(doc, "9. Technology Architecture (Non-AI Stack)", 1)
    add_para(
        doc,
        "This solution uses existing infrastructure and lightweight additions. No GPU, no ML pipeline, "
        "no speech-to-text service."
    )
    add_table(
        doc,
        ["Layer", "Function", "Recommended Tools (Examples)", "Cost Profile"],
        [
            ["Telephony / PBX", "Call routing, queue, recording, whisper/barge", "Asterisk, 3CX, Genesys, Twilio, RingCentral", "Existing license"],
            ["CRM / Ticketing", "Customer records, screen-pop, disposition logging", "Salesforce, HubSpot, Zendesk, Freshdesk", "Existing license"],
            ["Knowledge Base", "Playbooks, search, versioning", "Confluence, Notion, GitBook, internal wiki", "Low ($0–$500/mo)"],
            ["Wallboard", "Real-time metrics display", "PBX native, Geckoboard, custom Grafana dashboard", "Low ($0–$200/mo)"],
            ["Alerting", "Rule-based supervisor notifications", "Zapier, n8n, custom webhook, Slack/Teams bot", "Low ($0–$100/mo)"],
            ["QA Scorecard", "Call review and scoring", "Google Sheets, Airtable, or simple web app", "Low ($0)"],
            ["Post-Call Survey", "FCR/CSAT collection", "PBX IVR module, SurveyMonkey, Delighted", "Low ($0–$200/mo)"],
            ["CTI Integration", "Screen-pop on inbound call", "PBX webhook + CRM API (custom or connector)", "One-time setup"],
        ],
    )

    add_heading(doc, "9.1 Integration Architecture Diagram (Text)", 2)
    add_para(doc, "[Customer] → [IVR Menu] → [Skills Router] → [Agent Desktop]", bold=True)
    add_para(doc, "                              ↓                    ↓")
    add_para(doc, "                    [Self-Service Portal]   [CRM Screen-Pop]")
    add_para(doc, "                                              ↓")
    add_para(doc, "                                    [Playbook + Macros]")
    add_para(doc, "                                              ↓")
    add_para(doc, "                              [Help Button] → [Supervisor Alert]")
    add_para(doc, "                                              ↓")
    add_para(doc, "                                    [Wallboard + Whisper/Barge]")
    add_para(doc, "                                              ↓")
    add_para(doc, "                              [Wrap-Up Code] → [QA Sample + Survey]")

    # 10. Roadmap
    add_heading(doc, "10. 90-Day Implementation Roadmap", 1)
    add_table(
        doc,
        ["Phase", "Weeks", "Deliverables", "Owner", "Success Criteria"],
        [
            [
                "1. Foundation",
                "1–2",
                "Map top 20 call types; draft playbooks; define disposition codes",
                "Operations + Team Leads",
                "20 playbooks drafted; disposition taxonomy approved",
            ],
            [
                "2. Agent Tools",
                "3–4",
                "Publish knowledge base; deploy CTI screen-pop; create canned response macros",
                "IT + Operations",
                "Screen-pop live; agents trained on playbooks",
            ],
            [
                "3. Supervisor Center",
                "5–6",
                "Deploy wallboard; configure rule alerts; enable whisper/barge; launch help button",
                "IT + Supervisors",
                "Wallboard live; alerts firing; supervisor response < 60s",
            ],
            [
                "4. Quality Program",
                "7–8",
                "Finalize QA scorecard; begin weekly sampling; deploy post-call survey",
                "QA Team",
                "5 calls/agent/week reviewed; survey response rate > 20%",
            ],
            [
                "5. Customer Layer",
                "9–10",
                "Optimize IVR routing; publish self-service for top 5 drivers; enable callback",
                "Operations + IT",
                "Deflection rate measurable; abandon rate declining",
            ],
            [
                "6. Measure & Iterate",
                "11–12",
                "Baseline vs. current KPI comparison; refine playbooks; coach low-QA-score agents",
                "All stakeholders",
                "AHT, FCR, QA scores trending toward targets",
            ],
        ],
    )

    add_heading(doc, "10.1 Weekly Milestone Checklist", 2)
    milestones = {
        "Week 1": "Kickoff meeting; pull 90-day CDR data; identify top 20 call types by volume",
        "Week 2": "Complete playbook drafts; approve disposition code taxonomy",
        "Week 3": "Publish knowledge base; begin CTI integration development",
        "Week 4": "CTI screen-pop UAT; agent training session #1 (playbooks + macros)",
        "Week 5": "Wallboard deployment; configure first 5 alert rules",
        "Week 6": "Enable supervisor whisper/barge; launch help button; supervisor training",
        "Week 7": "QA scorecard finalized; select first week's random sample",
        "Week 8": "Post-call IVR survey live; first weekly QA report published",
        "Week 9": "IVR menu redesign based on disposition data; skills routing update",
        "Week 10": "Self-service pages published for top 5 call drivers",
        "Week 11": "KPI dashboard: compare baseline vs. current for AHT, FCR, abandon rate",
        "Week 12": "Executive review; document lessons learned; plan Phase 2 enhancements",
    }
    for week, task in milestones.items():
        add_bullet(doc, f"{week}: {task}")

    # 11. KPIs
    add_heading(doc, "11. KPIs, Targets & Measurement Framework", 1)
    add_table(
        doc,
        ["KPI", "Baseline (Set Week 1)", "30-Day Target", "90-Day Target", "Data Source"],
        [
            ["Average Handle Time", "Measure from CDR", "5% reduction", "15–25% reduction", "PBX CDR reports"],
            ["First Call Resolution", "Measure from survey + callbacks", "5% improvement", "10–20% improvement", "Post-call survey + repeat call analysis"],
            ["QA Score (avg)", "First sample week average", "Maintain or improve", "80+ average across team", "QA scorecard"],
            ["QA Agent Coverage", "0%", "80% of agents", "100% of agents weekly", "QA sampling log"],
            ["Queue Abandonment", "Measure from queue stats", "15% reduction", "30–50% reduction", "PBX queue reports"],
            ["Supervisor Response Time", "N/A", "< 120 seconds", "< 60 seconds", "Help button event log"],
            ["Self-Service Deflection", "0%", "5% of top-5 drivers", "15–25% of top-5 drivers", "IVR/portal analytics"],
            ["New Agent Time-to-Proficiency", "Current average", "10% improvement", "30–40% improvement", "Agent productivity reports"],
        ],
    )

    # 12. Governance
    add_heading(doc, "12. Governance & Human-in-the-Loop Controls", 1)
    add_para(
        doc,
        "Mirroring the hackathon's human-in-the-loop requirement, this solution ensures humans "
        "retain final authority on every customer interaction:"
    )
    add_bullet(doc, "Playbooks are authored and approved by subject-matter experts — not auto-generated")
    add_bullet(doc, "Agents choose which canned response to use — no auto-send")
    add_bullet(doc, "Supervisors decide when and how to intervene — alerts are recommendations, not automations")
    add_bullet(doc, "QA scores trigger coaching conversations — not automated disciplinary actions")
    add_bullet(doc, "Disposition codes are agent-selected — ensuring accountability and accuracy")
    add_bullet(doc, "Monthly playbook review cycle: operations team updates based on disposition trends and QA findings")

    # 13. Risk Assessment
    add_heading(doc, "13. Risk Assessment & Mitigation", 1)
    add_table(
        doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["Agent adoption resistance", "Medium", "High", "Involve agents in playbook creation; pilot with volunteer team first"],
            ["CTI integration delays", "Medium", "Medium", "Start with manual CRM lookup; automate in Phase 2 if needed"],
            ["Playbook content becomes stale", "High", "Medium", "Monthly review cadence; disposition data flags outdated playbooks"],
            ["Supervisor alert fatigue", "Medium", "Medium", "Tune alert thresholds after Week 6; prioritize high-severity only"],
            ["Low survey response rate", "Medium", "Low", "Keep survey to 1 question; offer IVR option immediately after call"],
            ["Insufficient QA staffing", "Low", "Medium", "5 calls/agent/week is ~1 FTE per 50 agents; plan capacity upfront"],
        ],
    )

    # 14. Budget
    add_heading(doc, "14. Budget & Resource Estimates", 1)
    add_table(
        doc,
        ["Resource", "Role", "Time Commitment (90 days)", "Cost Estimate"],
        [
            ["Project Lead", "Operations Manager", "25% FTE", "Internal"],
            ["IT Developer", "CTI integration + wallboard setup", "50% FTE for 6 weeks", "Internal or $5K–$15K contractor"],
            ["Team Leads (×3)", "Playbook authoring + agent training", "10% FTE each", "Internal"],
            ["QA Lead", "Scorecard design + sampling program", "50% FTE", "Internal"],
            ["Supervisors (×2)", "Alert tuning + coaching program", "5% FTE each", "Internal"],
            ["Knowledge base tool", "Confluence/Notion/GitBook", "—", "$0–$500/month"],
            ["Wallboard tool", "Geckoboard or Grafana", "—", "$0–$200/month"],
            ["Survey tool", "IVR module or Delighted", "—", "$0–$200/month"],
            ["TOTAL (excluding existing licenses)", "—", "—", "$5K–$20K + internal labor"],
        ],
    )

  # 15. Appendices
    doc.add_page_break()
    add_heading(doc, "15. Appendices", 1)

    add_heading(doc, "Appendix A: Sample Call Playbook — Billing Error", 2)
    add_para(doc, "Call Type: Billing Error Inquiry", bold=True)
    add_para(doc, "Trigger Phrases: \"wrong charge\", \"billing mistake\", \"overcharged\", \"invoice error\"")
    add_para(doc, "Target AHT: 6 minutes", bold=True)
    add_para(doc, "Resolution Steps:", bold=True)
    steps = [
        "Greet caller and verify identity (account number + security question)",
        "Pull up latest invoice in CRM (auto-loaded via screen-pop)",
        "Ask customer to identify the specific charge in question",
        "Compare charge against service records and contract terms",
        "If error confirmed: explain correction process; offer immediate credit",
        "If charge is valid: explain charge breakdown with approved wording",
        "Document resolution in CRM ticket",
        "Select disposition code: BILL_ERR (corrected) or BILL_VLD (explained)",
        "Close with: \"Is there anything else I can help you with today?\"",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")
    add_para(doc, "Escalation Criteria:", bold=True)
    add_bullet(doc, "Customer requests supervisor → click Help button")
    add_bullet(doc, "Credit amount exceeds agent authority ($50) → transfer to billing supervisor")
    add_bullet(doc, "Customer threatens legal action → escalate immediately")

    add_heading(doc, "Appendix B: QA Scorecard Template", 2)
    add_table(
        doc,
        ["Section", "Score (1–5)", "Notes"],
        [
            ["Greeting & Identification", "", ""],
            ["Active Listening", "", ""],
            ["Problem Diagnosis", "", ""],
            ["Resolution Steps", "", ""],
            ["Compliance Statements", "", ""],
            ["Closing & Confirmation", "", ""],
            ["TOTAL (/30 → convert to /100)", "", ""],
        ],
    )

    add_heading(doc, "Appendix C: Disposition Code Taxonomy (Starter Set)", 2)
    add_table(
        doc,
        ["Code", "Description", "Category"],
        [
            ["BILL_ERR", "Billing error — corrected", "Billing"],
            ["BILL_VLD", "Billing inquiry — charge valid", "Billing"],
            ["REFUND", "Refund processed or initiated", "Billing"],
            ["CANCEL", "Account cancellation processed", "Account"],
            ["PWD_RST", "Password reset completed", "Account"],
            ["OUTAGE", "Service outage reported/resolved", "Technical"],
            ["TECH_T1", "Technical issue — resolved Tier 1", "Technical"],
            ["TECH_T2", "Technical issue — escalated Tier 2", "Technical"],
            ["SHIP_DLY", "Shipping delay addressed", "Orders"],
            ["RETURN", "Return/RMA processed", "Orders"],
            ["UPGRADE", "Upgrade inquiry handled", "Sales"],
            ["COMP_AGT", "Complaint about agent", "Escalation"],
            ["COMP_SVC", "Complaint about service", "Escalation"],
            ["INFO", "General information provided", "General"],
            ["NO_RES", "Issue not resolved — callback scheduled", "General"],
        ],
    )

    add_heading(doc, "Appendix D: Supervisor Daily Checklist", 2)
    checklist = [
        "Review wallboard at shift start — confirm all agents logged in and queues healthy",
        "Respond to all help-button alerts within 60 seconds",
        "Monitor long-call alerts (> 12 min) — whisper-coach if agent is struggling",
        "Check queue wait every 30 minutes — trigger callback offer if > 3 min",
        "Review yesterday's abandon rate and disposition report",
        "Conduct 1 coaching session with lowest QA-score agent from prior week",
        "Verify all agents completed wrap-up codes before end of shift",
        "End-of-shift summary: log escalations, unresolved issues, and routing problems",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "Appendix E: Agent Desktop Quick-Reference Card", 2)
    add_para(doc, "During every call:", bold=True)
    add_bullet(doc, "Customer record auto-opens (screen-pop) — verify identity before making changes")
    add_bullet(doc, "Search playbooks by call type or keyword — follow steps in order")
    add_bullet(doc, "Use canned responses for approved wording — do not improvise compliance statements")
    add_bullet(doc, "Click \"Request Help\" if you need supervisor assistance — do not let calls go off-track")
    add_bullet(doc, "Select correct disposition code at wrap-up — mandatory on every call")
    add_bullet(doc, "Target: resolve on first contact; if not possible, schedule callback with specific date/time")

    add_heading(doc, "Appendix F: Comparison — AI Hackathon Solution vs. This Non-AI Solution", 2)
    add_table(
        doc,
        ["Capability", "AI Hackathon Approach", "This Non-AI Approach", "Outcome Parity"],
        [
            ["Mid-call knowledge help", "AI suggests articles from transcription", "Agent searches approved playbooks", "Yes — faster with good search"],
            ["Escalation detection", "AI sentiment analysis", "Rule-based duration/transfer alerts + help button", "Yes — more explainable"],
            ["Supervisor visibility", "AI dashboard with sentiment scores", "Wallboard + rule alerts", "Yes — deterministic"],
            ["Live coaching", "AI-generated coaching prompts", "Whisper/barge via PBX", "Yes — human judgment"],
            ["QA at scale", "AI transcribes and scores 100% of calls", "5 calls/agent/week + post-call survey", "Partial — survey covers 100%, deep review is sampled"],
            ["Intent detection", "NLP classification", "IVR routing + disposition codes", "Yes — with slight delay"],
            ["Implementation time", "3–6 months (ML pipeline)", "90 days", "Non-AI is faster"],
            ["Ongoing cost", "AI licensing + GPU + maintenance", "Minimal tooling + internal labor", "Non-AI is cheaper"],
            ["Explainability", "Black-box model predictions", "Fully transparent rules and playbooks", "Non-AI is superior"],
        ],
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Document —")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    document = build_document()
    output_path = "/workspace/Contact_Center_Non-AI_Solution.docx"
    document.save(output_path)
    print(f"Document saved to: {output_path}")
